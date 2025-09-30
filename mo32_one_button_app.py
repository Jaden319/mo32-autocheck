
import io, os, uuid, sqlite3, tempfile
from datetime import date, datetime, timedelta
import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image

APP_VER = "v8.5 Database added SQLite. "
st.set_page_config(page_title="MO32 Crane Compliance - Auto Check", layout="wide")

TODAY = date.today()
TODAY_STR = TODAY.strftime("%d-%m-%Y")  # DD-MM-YYYY
DATE_FORMATS = ("%d-%m-%Y","%d/%m/%Y","%Y-%m-%d")

# --- SQLite helper (uses a DB in the temp folder so it's writable on Streamlit Cloud) ---
TMP_ROOT = tempfile.gettempdir()
DB_PATH = os.path.join(TMP_ROOT, "inspections.db")
os.makedirs(TMP_ROOT, exist_ok=True)

def db_init():
    try:
        con = sqlite3.connect(DB_PATH, check_same_thread=False)
        con.execute("""
        CREATE TABLE IF NOT EXISTS inspections (
            id TEXT PRIMARY KEY,
            vessel TEXT,
            imo TEXT,
            created_at TEXT,
            case_dir TEXT,
            docx_path TEXT
        )
        """)
        con.commit()
        return con
    except Exception as e:
        st.warning(f"(DB unavailable: {e})")
        return None

DB = db_init()

def db_insert(vessel, imo, created_at, case_dir, docx_path):
    if not DB: return
    try:
        DB.execute("INSERT OR REPLACE INTO inspections (id, vessel, imo, created_at, case_dir, docx_path) VALUES (?, ?, ?, ?, ?, ?)",
                   (uuid.uuid4().hex, vessel or "", imo or "", created_at, case_dir or "", docx_path or ""))
        DB.commit()
    except Exception as e:
        st.warning(f"(DB insert skipped: {e})")

def db_search(vessel_like, imo_like):
    if not DB: return pd.DataFrame([], columns=["vessel","imo","created_at","case_dir","docx_path"])
    try:
        vl = f"%{vessel_like.strip()}%" if vessel_like else "%"
        il = f"%{imo_like.strip()}%" if imo_like else "%"
        df = pd.read_sql_query(
            "SELECT vessel, imo, created_at, case_dir, docx_path FROM inspections WHERE vessel LIKE ? AND imo LIKE ? ORDER BY created_at DESC",
            DB, params=(vl, il)
        )
        return df
    except Exception as e:
        st.warning(f"(DB search skipped: {e})")
        return pd.DataFrame([], columns=["vessel","imo","created_at","case_dir","docx_path"])

def asciiize(s):
    if s is None: return ""
    trans = {"–":"-","—":"-","-":"-","•":"*","·":"*","“":'"',"”":'"',"‘":"'", "’":"'", "…":"...", "°":" deg ","×":"x","✓":"OK","\u00a0":" "}
    out = str(s)
    for k,v in trans.items(): out = out.replace(k,v)
    try:
        out.encode("latin-1"); return out
    except UnicodeEncodeError:
        return out.encode("latin-1","ignore").decode("latin-1")

CHECK_COLUMNS = [
    "Crane #","Vessel Name","IMO","Crane Make/Model","Serial Number","SWL (t)",
    "Install/Commission Date","Last 5-year Proof Test Date","Last Annual Thorough Exam Date",
    "Annual Exam By (Competent/Responsible Person)","Certificate of Test # (AMSA 365/642/etc)",
    "Certificate Current? (Y/N)","Register of MHE Onboard? (Y/N)","Pre-use Visual Exam OK? (Y/N)",
    "Rigging Plan/Drawings Onboard? (Y/N)","Controls layout labelled & accessible? (Y/N)",
    "Limit switches operational? (Y/N)","Brakes operational? (Y/N)","Operator visibility adequate? (Y/N)",
    "Visibility: Shift (Day/Evening/Night)","Visibility: Weather conditions",
    "Weather protection at winch/controls? (Y/N)","Access/escape to cabin compliant? (Y/N)","Notes / Defects",
    "Loose Gear: Hook/Block Serial Number","Loose Gear: Hook SWL (t)","Loose Gear: Certificate Number",
    "Loose Gear: Last Inspection/Proof Date","Loose Gear: Notes"
]

GUIDANCE = [
    ("5-year proof load test interval", "MO32 Sch.3 2(2)(a)"),
    ("12-month thorough exam interval", "MO32 Sch.3 2(2)(b), 2(5)"),
    ("Certificates current / approved forms", "MO32 s.23"),
    ("Register of MHE kept onboard (maintenance & repair log)", "MO32 s.25"),
    ("Pre-use visual exam before operation", "MO32 s.22(2)(c)"),
    ("Rigging plan/drawings onboard", "MO32 Sch.6 Div.2 cl.1"),
    ("Controls & brakes (incl. limit switches)", "MO32 Sch.6 Div.3; Sch.3 cl.4(3)"),
    ("Weather protection at controls", "MO32 Sch.6 cl.19"),
    ("Loose Gear inspection & traceability", "Good practice; tie to crane SWL"),
]

RISK_KEYWORDS = {
  "Controls layout labelled & accessible? (Y/N)": ["sloppy","sticky","not labelled","label missing","hard to reach","unresponsive","stiff"],
  "Brakes operational? (Y/N)": ["drifts","creeps","won't hold","wont hold","weak brake","fade","brake fade","slips"],
  "Limit switches operational? (Y/N)": ["overtravel","limit not working","no cutout","failsafe","upper limit","lower limit"],
  "Operator visibility adequate? (Y/N)": ["blind spot","obstructed","poor visibility","camera not working","wiper","dirty screen"],
  "Weather protection at winch/controls? (Y/N)": ["no canopy","exposed","water ingress","leaks","rain on controls"],
  "Access/escape to cabin compliant? (Y/N)": ["ladder loose","handrail","blocked","trip hazard","missing step"],
  "Rigging Plan/Drawings Onboard? (Y/N)": ["no plan","drawing missing","outdated plan","no rigging plan"],
  "Certificate Current? (Y/N)": ["expired","out of date","no certificate","missing cert"],
  "Register of MHE Onboard? (Y/N)": ["no register","not onboard","out of date","not updated"],
}
WEATHER_BAD = ["rain","raining","storm","storming","hail","fog","mist","spray","squall","gust","windy","dust","smoke","night","dark","low light","glare"]

def safe_text(v):
    if v is None: return ""
    try:
        import math
        if isinstance(v, float) and (v!=v): return ""
    except Exception: pass
    s = str(v)
    return "" if s.strip().lower() in ("nan","none") else s

def parse_date(v):
    if not v: return None
    if isinstance(v,(pd.Timestamp,datetime)):
        return v.date() if hasattr(v,"date") else v
    s = str(v).strip()
    for fmt in DATE_FORMATS:
        try: return datetime.strptime(s,fmt).date()
        except Exception: pass
    try:
        return pd.to_datetime(float(s), unit="D", origin="1899-12-30").date()
    except Exception:
        return None

def to_float(v):
    try: return float(str(v).strip())
    except Exception: return None

def yn(v): return str(v).strip().upper()=="Y"
def days_since(d): return None if not d else (TODAY-d).days
def days_left_since(d, interval):
    if not d: return None
    due = d + timedelta(days=interval)
    return (due - TODAY).days

def contradiction_notes_check(row):
    notes = (safe_text(row.get("Notes / Defects"))+" "+safe_text(row.get("Loose Gear: Notes"))+" "+safe_text(row.get("Visibility: Weather conditions"))).lower()
    findings = []
    for field, words in RISK_KEYWORDS.items():
        tick = str(row.get(field,"")).strip().upper()
        if not notes or tick not in ("Y","N"): continue
        if tick=="Y" and any(w in notes for w in words):
            hits = [w for w in words if w in notes][:3]
            findings.append(f"{field}: Y but notes mention {', '.join(hits)}")
        if tick=="N" and any(x in notes for x in ["ok now","temporary fix","workaround","still usable"]):
            findings.append(f"{field}: N but notes imply workaround/operation")
    vis = str(row.get("Operator visibility adequate? (Y/N)","")).strip().upper()
    shift = safe_text(row.get("Visibility: Shift (Day/Evening/Night)")).lower()
    if vis=="Y" and (shift=="night" or any(w in notes for w in WEATHER_BAD)):
        findings.append("Visibility = Y but conditions indicate low visibility (night/adverse weather).")
    return findings

def evidence_prompts(row):
    prompts = []
    if yn(row.get("Certificate Current? (Y/N)")) and not safe_text(row.get("Certificate of Test # (AMSA 365/642/etc)")).strip():
        prompts.append("Certificate marked current but certificate number is blank - add ID/photo.")
    if yn(row.get("Register of MHE Onboard? (Y/N)")) and not safe_text(row.get("Annual Exam By (Competent/Responsible Person)")).strip():
        prompts.append("Register marked onboard - add last entry details/competent or responsible person.")
    if yn(row.get("Rigging Plan/Drawings Onboard? (Y/N)")):
        notes = safe_text(row.get("Notes / Defects")).lower()
        if "plan" in notes and "rev" not in notes:
            prompts.append("Rigging plan onboard - add drawing ID/revision/date in notes.")
    if yn(row.get("Certificate Current? (Y/N)")) and not safe_text(row.get("Loose Gear: Certificate Number")).strip():
        prompts.append("Main certificate current but loose gear cert # blank - add accessory cert reference.")
    return prompts

def evaluate_row(row):
    issues, attention, due_soon = [], [], []
    d5 = parse_date(row.get("Last 5-year Proof Test Date"))
    if not d5 or days_since(d5) > 5*365:
        issues.append("Overdue/missing 5-year proof test (MO32 Sch.3 2(2)(a)).")
    else:
        left = days_left_since(d5, 5*365)
        if left is not None and left <= 90:
            due_soon.append(f"5-year proof test due in {left} days.")
    d12 = parse_date(row.get("Last Annual Thorough Exam Date"))
    if not d12 or days_since(d12) > 365+31:
        issues.append("Overdue/missing annual thorough exam (MO32 Sch.3 2(2)(b), 2(5)).")
    else:
        left = days_left_since(d12, 365)
        if left is not None and left <= 30:
            due_soon.append(f"Annual thorough exam due in {left} days.")
    for field, ref in [
        ("Certificate Current? (Y/N)", "s.23"),
        ("Register of MHE Onboard? (Y/N)", "s.25"),
        ("Pre-use Visual Exam OK? (Y/N)", "s.22(2)(c)"),
        ("Rigging Plan/Drawings Onboard? (Y/N)", "Sch.6 Div.2 cl.1"),
        ("Limit switches operational? (Y/N)", "Sch.3 cl.4(3)"),
        ("Brakes operational? (Y/N)", "Sch.6 Div.3"),
        ("Controls layout labelled & accessible? (Y/N)", "Sch.6 Div.3"),
        ("Operator visibility adequate? (Y/N)", "Sch.6 vis."),
        ("Weather protection at winch/controls? (Y/N)", "Sch.6 cl.19"),
        ("Access/escape to cabin compliant? (Y/N)", "Sch.6 access"),
    ]:
        tick = str(row.get(field,"")).strip().upper()
        if tick not in ("Y","N"):
            attention.append(f"{field} not answered (tick Y or N).")
        elif tick != "Y":
            issues.append(f"{field.replace(' (Y/N)','')}: NOT OK (MO32 {ref}).")
    shift = safe_text(row.get("Visibility: Shift (Day/Evening/Night)")).strip()
    weather = safe_text(row.get("Visibility: Weather conditions")).strip()
    if not shift:
        attention.append("Shift/Lighting not selected (Day/Evening/Night).")
    if not weather:
        attention.append("Weather conditions box is empty (e.g., Raining / Clear / Fog).")
    crane_swl = to_float(row.get("SWL (t)"))
    hook_swl = to_float(row.get("Loose Gear: Hook SWL (t)"))
    if hook_swl is not None and crane_swl is not None and hook_swl > crane_swl:
        issues.append(f"Loose gear SWL ({hook_swl} t) exceeds crane SWL ({crane_swl} t) - mismatch.")
    lg_date = parse_date(row.get("Loose Gear: Last Inspection/Proof Date"))
    if not lg_date or days_since(lg_date) > 365:
        issues.append("Loose gear last inspection/proof >12 months or missing - not compliant.")
    else:
        left = days_left_since(lg_date, 365)
        if left is not None and left <= 30:
            due_soon.append(f"Loose gear inspection due in {left} days.")
    if yn(row.get("Certificate Current? (Y/N)")) and not safe_text(row.get("Loose Gear: Certificate Number")).strip():
        attention.append("Loose gear certificate number blank while main cert is current - add accessory cert ref.")
    contradictions = contradiction_notes_check(row)
    if contradictions:
        attention.extend([f"Notes contradict ticks: {c}" for c in contradictions])
    prompts = evidence_prompts(row)
    if prompts:
        attention.extend(prompts)
    status = "PASS"
    if issues:
        status = "FAIL"
    elif attention or due_soon:
        status = "ATTENTION"
    return status, issues, attention, due_soon

def evaluate(df):
    missing = [c for c in CHECK_COLUMNS if c not in df.columns]
    if missing:
        raise ValueError("Form columns missing: " + ", ".join(missing))
    rows = []
    for _, row in df.iterrows():
        status, issues, attention, due_soon = evaluate_row(row)
        rows.append({
            "Crane #": row.get("Crane #"),
            "Vessel Name": row.get("Vessel Name"),
            "IMO": row.get("IMO"),
            "Serial Number": row.get("Serial Number"),
            "SWL (t)": row.get("SWL (t)"),
            "Shift": row.get("Visibility: Shift (Day/Evening/Night)"),
            "Weather": row.get("Visibility: Weather conditions"),
            "Loose Gear Serial": row.get("Loose Gear: Hook/Block Serial Number"),
            "Loose Gear SWL (t)": row.get("Loose Gear: Hook SWL (t)"),
            "Status": status,
            "Issues (FAIL)": "; ".join(issues) if issues else "",
            "Attention (notes/evidence)": "; ".join(attention) if attention else "",
            "Due soon": "; ".join(due_soon) if due_soon else "",
        })
    return pd.DataFrame(rows)

# --- Upload helpers ---
def ensure_jpeg(data_bytes):
    try:
        im = Image.open(io.BytesIO(data_bytes))
        rgb = im.convert("RGB")
        out = io.BytesIO()
        rgb.save(out, format="JPEG", quality=85)
        return out.getvalue()
    except Exception:
        return data_bytes

def _shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color.replace('#',''))
    tcPr.append(shd)

def _add_heading(doc, text, size=16):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    return p

def build_docx(results_df, df_original, photos_map, photos_loose_map, out_path=None):
    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    # Cover header
    t = doc.add_table(rows=1, cols=2)
    left, right = t.rows[0].cells
    _shade_cell(left,  "1F4E79")
    _shade_cell(right, "1F4E79")
    p1 = left.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p1.add_run("Crane Compliance Inspection Report")
    r1.font.color.rgb = RGBColor(255,255,255); r1.bold = True; r1.font.size = Pt(22)
    p2 = right.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run("Your Logo")
    r2.font.color.rgb = RGBColor(255,255,255); r2.bold = True; r2.font.size = Pt(14)

    doc.add_paragraph("")
    vessel_name = asciiize(df_original['Vessel Name'].iloc[0]) if 'Vessel Name' in df_original.columns and len(df_original) else ""
    imo_code = asciiize(df_original['IMO'].iloc[0]) if 'IMO' in df_original.columns and len(df_original) else ""
    inspector = asciiize(safe_text(df_original.get("Inspector / Role", pd.Series([''])).iloc[0])) if "Inspector / Role" in df_original.columns and len(df_original) else ""
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(6)
    meta.add_run("Marine Orders 32 (MO32) – Stevedore AutoCheck\n").bold = True
    meta.add_run(f"Vessel: {vessel_name}    IMO: {imo_code}\n")
    meta.add_run(f"Date: {TODAY_STR}    Inspector: {inspector}")

    doc.add_page_break()

    # Executive Summary
    _add_heading(doc, "Executive Summary", 18)
    total = len(results_df)
    c_pass = int((results_df["Status"]=="PASS").sum()) if "Status" in results_df else 0
    c_attn = int((results_df["Status"]=="ATTENTION").sum()) if "Status" in results_df else 0
    c_fail = int((results_df["Status"]=="FAIL").sum()) if "Status" in results_df else 0

    kpi = doc.add_table(rows=1, cols=3)
    kpi.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,(lab,val,color) in enumerate([("PASS",c_pass,"92D050"),("ATTENTION",c_attn,"FFC000"),("FAIL",c_fail,"FF0000")]):
        cell = kpi.rows[0].cells[i]
        _shade_cell(cell, color)
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{lab}\n{val} of {total}")
        run.font.size = Pt(14); run.bold = True
        run.font.color.rgb = RGBColor(0,0,0) if lab=="ATTENTION" else RGBColor(255,255,255)

    doc.add_paragraph("")

    headers = ["Crane #","Serial #","SWL (t)","Shift","Weather","Status"]
    tbl = doc.add_table(rows=1, cols=len(headers), style="Table Grid")
    for i,h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        try:
            cell.paragraphs[0].runs[0].bold = True
        except Exception:
            pass
        _shade_cell(cell, "D9D9D9")
    for _, r in results_df.iterrows():
        row = tbl.add_row().cells
        vals = [
            str(r.get("Crane #","")),
            asciiize(safe_text(r.get("Serial Number",""))),
            safe_text(r.get("SWL (t)","")),
            asciiize(safe_text(r.get("Shift",""))),
            asciiize(safe_text(r.get("Weather",""))),
            safe_text(r.get("Status","")),
        ]
        for i,v in enumerate(vals): row[i].text = v
        scell = row[-1]
        if vals[-1]=="PASS": _shade_cell(scell,"92D050")
        elif vals[-1]=="ATTENTION": _shade_cell(scell,"FFC000")
        elif vals[-1]=="FAIL": _shade_cell(scell,"FF0000")

    doc.add_page_break()

    # Per-crane sections
    for crane_no in [1,2,3,4]:
        sub = df_original[df_original["Crane #"]==crane_no] if "Crane #" in df_original.columns else df_original.iloc[0:0]
        imgs = photos_map.get(crane_no) or []
        imgs_lg = photos_loose_map.get(crane_no) or []
        if sub.empty and not imgs and not imgs_lg: continue

        _add_heading(doc, f"Crane {crane_no}", 16)

        if not sub.empty:
            row = sub.iloc[0]
            key_tbl = doc.add_table(rows=0, cols=2, style="Table Grid")
            def add_row(label, value):
                r = key_tbl.add_row().cells
                r[0].text = label; 
                try: r[0].paragraphs[0].runs[0].bold = True
                except Exception: pass
                _shade_cell(r[0], "EEEEEE")
                r[1].text = asciiize(safe_text(value))
            add_row("Make/Model", row.get("Crane Make/Model"))
            add_row("Serial Number", row.get("Serial Number"))
            add_row("SWL (t)", row.get("SWL (t)"))
            add_row("Install/Commission Date", row.get("Install/Commission Date"))
            add_row("Last 5-year Proof Test Date", row.get("Last 5-year Proof Test Date"))
            add_row("Last Annual Thorough Exam Date", row.get("Last Annual Thorough Exam Date"))
            add_row("Exam By (Competent/Responsible Person)", row.get("Annual Exam By (Competent/Responsible Person)"))
            add_row("Certificate of Test #", row.get("Certificate of Test # (AMSA 365/642/etc)"))
            add_row("Shift / Weather", f"{row.get('Visibility: Shift (Day/Evening/Night)')} / {row.get('Visibility: Weather conditions')}")

            doc.add_paragraph("")

            ticks = [
                ("Certificate Current?", row.get("Certificate Current? (Y/N)")),
                ("Register of MHE Onboard?", row.get("Register of MHE Onboard? (Y/N)")),
                ("Pre-use Visual Exam OK?", row.get("Pre-use Visual Exam OK? (Y/N)")),
                ("Rigging Plan/Drawings Onboard?", row.get("Rigging Plan/Drawings Onboard? (Y/N)")),
                ("Controls labelled & accessible?", row.get("Controls layout labelled & accessible? (Y/N)")),
                ("Limit switches operational?", row.get("Limit switches operational? (Y/N)")),
                ("Brakes operational?", row.get("Brakes operational? (Y/N)")),
                ("Operator visibility adequate?", row.get("Operator visibility adequate? (Y/N)")),
                ("Weather protection at controls?", row.get("Weather protection at winch/controls? (Y/N)")),
                ("Access/escape compliant?", row.get("Access/escape to cabin compliant? (Y/N)")),
            ]
            tick_tbl = doc.add_table(rows=1, cols=2, style="Table Grid")
            tick_tbl.rows[0].cells[0].text = "Item"; 
            try: tick_tbl.rows[0].cells[0].paragraphs[0].runs[0].bold = True
            except Exception: pass
            tick_tbl.rows[0].cells[1].text = "Y/N"; 
            try: tick_tbl.rows[0].cells[1].paragraphs[0].runs[0].bold = True
            except Exception: pass
            _shade_cell(tick_tbl.rows[0].cells[0], "D9D9D9"); _shade_cell(tick_tbl.rows[0].cells[1], "D9D9D9")
            for lab,val in ticks:
                rr = tick_tbl.add_row().cells
                rr[0].text = lab
                rr[1].text = safe_text(val)

            notes = safe_text(row.get("Notes / Defects"))
            lg_notes = safe_text(row.get("Loose Gear: Notes"))
            if notes or lg_notes:
                _add_heading(doc, "Notes / Defects", 14)
                if notes: doc.add_paragraph(asciiize(notes), style="List Paragraph")
                if lg_notes: doc.add_paragraph(asciiize(lg_notes), style="List Paragraph")

            rr = None
            try:
                rr = results_df[results_df["Crane #"]==crane_no]
            except Exception:
                pass
            if rr is not None and not rr.empty:
                issues = safe_text(rr.iloc[0]["Issues (FAIL)"])
                attn   = safe_text(rr.iloc[0]["Attention (notes/evidence)"])
                due    = safe_text(rr.iloc[0]["Due soon"])

                _add_heading(doc, "Compliance Findings", 14)
                box = doc.add_table(rows=3, cols=1, style="Table Grid")
                cell = box.rows[0].cells[0]; _shade_cell(cell, "FF0000")
                cell.paragraphs[0].add_run("ISSUES (FAIL): ").bold = True
                cell.paragraphs[0].add_run(asciiize(issues) if issues else "None recorded.")
                cell = box.rows[1].cells[0]; _shade_cell(cell, "FFC000")
                cell.paragraphs[0].add_run("ATTENTION: ").bold = True
                cell.paragraphs[0].add_run(asciiize(attn) if attn else "None recorded.")
                cell = box.rows[2].cells[0]; _shade_cell(cell, "D9E1F2")
                cell.paragraphs[0].add_run("DUE SOON: ").bold = True
                cell.paragraphs[0].add_run(asciiize(due) if due else "None.")

        def add_gallery(title, blobs):
            if not blobs: return
            _add_heading(doc, title, 12)
            t = doc.add_table(rows=0, cols=2, style="Table Grid")
            count = 0
            for i, data in enumerate(blobs[:8]):
                if count % 2 == 0:
                    row = t.add_row().cells
                try:
                    bio = io.BytesIO(data); bio.seek(0)
                    p = row[count % 2].paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(bio, width=Inches(3.0))
                    cap = row[count % 2].add_paragraph(f"{title} {i+1}")
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    row[count % 2].text = f"(Could not embed image: {e})"
                count += 1

        add_gallery("Crane Photos", imgs)
        add_gallery("Loose Gear Photos", imgs_lg)

        doc.add_page_break()

    _add_heading(doc, "Applicable MO32 Clauses & Guidance", 16)
    for text,ref in GUIDANCE:
        doc.add_paragraph(f"{text} — {ref}", style="List Bullet")

    if out_path:
        doc.save(out_path); return None
    buff = io.BytesIO(); doc.save(buff); buff.seek(0); return buff

def save_case(results_df, df_original, photos_map, photos_loose_map):
    # Use temp dir to avoid permission issues and file-collision on Streamlit Cloud
    base = os.path.join(TMP_ROOT, "mo32_cases")
    # If something named base exists but is NOT a dir, remove and recreate
    if os.path.exists(base) and not os.path.isdir(base):
        try: os.remove(base)
        except Exception: pass
    os.makedirs(base, exist_ok=True)

    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    uniq = uuid.uuid4().hex[:6]
    case_dir = os.path.join(base, f"case_{stamp}_{uniq}")
    os.makedirs(case_dir, exist_ok=False)

    df_original.to_csv(os.path.join(case_dir,"inputs.csv"), index=False)
    results_df.to_csv(os.path.join(case_dir,"results.csv"), index=False)
    for label, mapping in [("crane",photos_map), ("loose",photos_loose_map)]:
        for k, blobs in mapping.items():
            if not blobs: continue
            pdir = os.path.join(case_dir, f"crane_{k}_{label}_photos"); os.makedirs(pdir, exist_ok=True)
            for i,data in enumerate(blobs):
                with open(os.path.join(pdir, f"photo_{i+1}.jpg"), "wb") as imgf:
                    imgf.write(data)

    build_docx(results_df, df_original, photos_map, photos_loose_map, out_path=os.path.join(case_dir,"MO32_Crane_Compliance_Report.docx"))
    return case_dir

# -------------------------
# Page: Vessel Inspection
# -------------------------
def page_inspection():
    st.title("Vessel Inspection")
    st.caption(APP_VER + " – More version to be updated with. good for testing.")

    with st.sidebar:
        st.header("Options")
        st.caption("DOCX only (photos embedded; DD-MM-YYYY)")
        st.subheader("CSV (optional)")
        csv_up = st.file_uploader("Import CSV (optional)", type=["csv"], key="u_csv")
        if csv_up is not None:
            try:
                df_loaded = pd.read_csv(csv_up)
                st.success("CSV imported. Use 'Evaluate imported CSV' to generate DOCX.")
                if st.button("Evaluate imported CSV", key="btn_eval_csv"):
                    out_df = evaluate(df_loaded)
                    st.subheader("Results (PASS/ATTENTION/FAIL) — Imported CSV")
                    st.dataframe(out_df, use_container_width=True)
                    docx_io = build_docx(out_df, df_loaded, {1:[],2:[],3:[],4:[]}, {1:[],2:[],3:[],4:[]})
                    st.download_button("Download Word report (.docx)", docx_io.getvalue(), file_name="MO32_Crane_Compliance_Report.docx", key="dl_docx_csv")
            except Exception as e:
                st.error(f"Error reading CSV: {e}")
        if st.button("Download blank CSV", key="btn_blankcsv"):
            df_blank = pd.DataFrame([[i+1] + [""]*(len(CHECK_COLUMNS)-1) for i in range(4)], columns=CHECK_COLUMNS)
            st.download_button("Save blank CSV", df_blank.to_csv(index=False).encode("utf-8"), file_name="Crane_Compliance_MO32_Blank.csv", key="dl_blankcsv")

    st.markdown("### Job/Vessel info")
    colv1, colv2, colv3 = st.columns(3)
    with colv1: vessel = st.text_input("Vessel Name","", key="vessel")
    with colv2: imo = st.text_input("IMO","", key="imo")
    with colv3: operator = st.text_input("Inspector / Role","", key="operator")

    st.markdown("### Crane checks")
    crane_data = []; photos_map = {}; photos_loose_map = {}
    for n in [1,2,3,4]:
        with st.expander(f"Crane {n}", expanded=(n==1)):
            c1,c2,c3 = st.columns(3)
            with c1:
                make_model = st.text_input(f"Crane {n} Make/Model", key=f"mm{n}")
                serial = st.text_input(f"Crane {n} Serial Number", key=f"sn{n}")
            with c2:
                swl = st.text_input(f"Crane {n} SWL (t)", key=f"swl{n}")
                install = st.text_input(f"Install/Commission Date (DD-MM-YYYY)", key=f"inst{n}")
            with c3:
                proof = st.text_input("Last 5-year Proof Test Date (DD-MM-YYYY)", key=f"p5{n}")
                annual = st.text_input("Last Annual Thorough Exam Date (DD-MM-YYYY)", key=f"a12{n}")
            c4,c5 = st.columns(2)
            with c4:
                annual_by = st.text_input("Annual Exam By (Competent/Responsible Person)", key=f"by{n}")
                cert_no = st.text_input("Certificate of Test # (AMSA 365/642/etc)", key=f"cert{n}")
            with c5:
                st.markdown("**Y/N items** (tick Y if compliant)")
                y_cert = st.selectbox("Certificate Current? (AMSA 365/642 form)", ["","Y","N"], key=f"yc{n}")
                y_reg  = st.selectbox("Register of MHE Onboard? (Maintenance & repair log)", ["","Y","N"], key=f"yr{n}")
                y_pre  = st.selectbox("Pre-use Visual Exam OK? (before operation)", ["","Y","N"], key=f"yp{n}")
                y_plan = st.selectbox("Rigging Plan/Drawings Onboard? (latest revision available)", ["","Y","N"], key=f"ypl{n}")
                y_ctrl = st.selectbox("Controls labelled & accessible? (labels present, reachable)", ["","Y","N"], key=f"yct{n}")
                y_lim  = st.selectbox("Limit switches operational?", ["","Y","N"], key=f"yl{n}")
                y_brk  = st.selectbox("Brakes operational?", ["","Y","N"], key=f"yb{n}")
                y_vis  = st.selectbox("Operator visibility adequate? (consider lighting & weather)", ["","Y","N"], key=f"yv{n}")
                y_wth  = st.selectbox("Weather protection at controls? (canopy/cover, no ingress)", ["","Y","N"], key=f"yw{n}")
                y_acc  = st.selectbox("Access/escape to cabin compliant? (ladder/handrails clear)", ["","Y","N"], key=f"ya{n}")
            w1,w2 = st.columns([1,2])
            with w1:
                shift = st.selectbox("Shift/Lighting", ["","Day","Evening","Night"], key=f"shift{n}")
            with w2:
                wx = st.text_input("Weather conditions (e.g., Raining, Storming, Fog, Clear)", key=f"wx{n}")
            notes = st.text_area("Notes / Defects", key=f"notes{n}", height=100)

            up_crane = st.file_uploader(f"Crane {n} photos (JPG/PNG/HEIC; up to 8)", type=["jpg","jpeg","png","heic","heif"], accept_multiple_files=True, key=f"photos{n}")
            photos_map[n] = [ensure_jpeg(f.getvalue()) for f in (up_crane or [])]

            st.markdown("#### Loose Gear (hook/block)")
            lg1, lg2, lg3 = st.columns(3)
            with lg1:
                lg_serial = st.text_input("Hook/Block Serial Number", key=f"lgsn{n}")
                lg_cert   = st.text_input("Certificate Number", key=f"lgcert{n}")
            with lg2:
                lg_swl = st.text_input("Hook SWL (t)", key=f"lgswl{n}")
                lg_date = st.text_input("Last Inspection/Proof Date (DD-MM-YYYY)", key=f"lgdate{n}")
            with lg3:
                lg_notes = st.text_area("Loose Gear Notes", key=f"lgnotes{n}", height=80)

            up_loose = st.file_uploader(f"Crane {n} loose gear photos (JPG/PNG/HEIC; up to 6)", type=["jpg","jpeg","png","heic","heif"], accept_multiple_files=True, key=f"photos_loose{n}")
            photos_loose_map[n] = [ensure_jpeg(f.getvalue()) for f in (up_loose or [])]

            crane_data.append({
                "Crane #": n, "Vessel Name": vessel, "IMO": imo,
                "Crane Make/Model": make_model, "Serial Number": serial, "SWL (t)": swl,
                "Install/Commission Date": install, "Last 5-year Proof Test Date": proof, "Last Annual Thorough Exam Date": annual,
                "Annual Exam By (Competent/Responsible Person)": annual_by, "Certificate of Test # (AMSA 365/642/etc)": cert_no,
                "Certificate Current? (Y/N)": y_cert, "Register of MHE Onboard? (Y/N)": y_reg, "Pre-use Visual Exam OK? (Y/N)": y_pre,
                "Rigging Plan/Drawings Onboard? (Y/N)": y_plan, "Controls layout labelled & accessible? (Y/N)": y_ctrl,
                "Limit switches operational? (Y/N)": y_lim, "Brakes operational? (Y/N)": y_brk, "Operator visibility adequate? (Y/N)": y_vis,
                "Visibility: Shift (Day/Evening/Night)": shift, "Visibility: Weather conditions": wx,
                "Weather protection at winch/controls? (Y/N)": y_wth, "Access/escape to cabin compliant? (Y/N)": y_acc, "Notes / Defects": notes,
                "Loose Gear: Hook/Block Serial Number": lg_serial, "Loose Gear: Hook SWL (t)": lg_swl, "Loose Gear: Certificate Number": lg_cert,
                "Loose Gear: Last Inspection/Proof Date": lg_date, "Loose Gear: Notes": lg_notes
            })

    st.divider()
    left, mid, right = st.columns([2,1,1])
    eval_clicked = left.button("Evaluate & Generate Report (DOCX)", type="primary", use_container_width=True, key="btn_eval")
    csv_clicked  = mid.button("Download current inputs as CSV", use_container_width=True, key="btn_csv")
    demo_clicked = right.button("Try demo (sample data)", use_container_width=True, key="btn_demo")

    if csv_clicked:
        df_now = pd.DataFrame(crane_data, columns=CHECK_COLUMNS)
        st.download_button("Save this CSV now", df_now.to_csv(index=False).encode("utf-8"), file_name="Crane_Compliance_MO32_Current.csv", key="dl_currentcsv")

    if eval_clicked:
        df_input = pd.DataFrame(crane_data, columns=CHECK_COLUMNS)
        try:
            out_df = evaluate(df_input)
            st.subheader("Results (PASS/ATTENTION/FAIL)")
            st.dataframe(out_df, use_container_width=True)
            st.success("Evaluation complete. Download your DOCX report below.")
            docx_io = build_docx(out_df, df_input, photos_map, photos_loose_map)
            st.download_button("Download Word report (.docx)", docx_io.getvalue(), file_name="MO32_Crane_Compliance_Report.docx", key="dl_docx_real")
            case_dir = save_case(out_df, df_input, photos_map, photos_loose_map)
            docx_path = os.path.join(case_dir, "MO32_Crane_Compliance_Report.docx")

            # Insert into DB
            vessel_val = str(df_input.get("Vessel Name").iloc[0]) if "Vessel Name" in df_input.columns and len(df_input) else ""
            imo_val = str(df_input.get("IMO").iloc[0]) if "IMO" in df_input.columns and len(df_input) else ""
            created_at = datetime.now().isoformat(timespec="seconds")
            db_insert(vessel_val, imo_val, created_at, case_dir, docx_path)

            st.info(f"Saved a copy of this submission to: {case_dir}")
        except Exception as e:
            st.error(f"Error during evaluation: {e}")

    if demo_clicked:
        demo_df = pd.DataFrame([
            [1,"MV Example","9526722","NMF DKII","CRN-123","45","10-05-2019","08-05-2022","01-06-2025","Joe Bloggs (Comp)","AMSA365-111","Y","Y","Y","Y","Y","Y","Y","Y","Night","Raining",
             "Y","Y","Controls slightly sloppy at low speed",
             "LG-001","40","LGCERT-789","15-06-2025","hook latch slightly bent"]
        ], columns=CHECK_COLUMNS)
        try:
            out_df = evaluate(demo_df)
            st.subheader("Results (PASS/ATTENTION/FAIL) - Demo")
            st.dataframe(out_df, use_container_width=True)
            docx_io = build_docx(out_df, demo_df, {1:[],2:[],3:[],4:[]}, {1:[],2:[],3:[],4:[]})
            st.download_button("Download Word report (.docx) - Demo", docx_io.getvalue(), file_name="MO32_Crane_Compliance_Report_DEMO.docx", key="dl_docx_demo")
        except Exception as e:
            st.error(f"Error during demo evaluation: {e}")

# -------------------------
# Page: Search Vessels
# -------------------------
def page_search():
    st.title("Search Vessels")
    st.caption("Search DB (SQLite) for Reports.")

    q_name = st.text_input("Vessel Name (partial ok)", key="q_vessel").strip()
    q_imo  = st.text_input("IMO Number (exact or partial)", key="q_imo").strip()

    # --- DB results
    st.markdown("#### Database results (SQLite)")
    df_db = db_search(q_name, q_imo)
    if df_db is not None and len(df_db):
        for i, r in df_db.iterrows():
            c1, c2, c3, c4 = st.columns([3,2,2,2])
            c1.markdown(f"**{r['vessel'] or '(No Vessel Name)'}**  \nIMO: {r['imo'] or '-'}")
            c2.markdown(f"**Created:** {r['created_at']}")
            c3.markdown(f"**Case dir:** {os.path.basename(r['case_dir'] or '')}")
            docx_path = r.get("docx_path")
            if docx_path and os.path.isfile(docx_path):
                try:
                    data = open(docx_path, "rb").read()
                    c4.download_button("Download DOCX", data, file_name=f"{os.path.basename(r['case_dir']) or 'report'}.docx", key=f"dl_db_{i}")
                except Exception as e:
                    c4.write(f"(DOCX not readable: {e})")
            else:
                c4.write("(No DOCX)")
            st.divider()
    else:
        st.info("No matches in database.")

    # --- Legacy folder scan in temp
    st.markdown("#### Saved cases in temp folder")
    base = os.path.join(TMP_ROOT, "mo32_cases")
    if not os.path.isdir(base):
        st.info("No saved case folders yet.")
        return

    rows = []
    for d in sorted(os.listdir(base)):
        case_dir = os.path.join(base, d)
        if not os.path.isdir(case_dir): continue
        inputs = os.path.join(case_dir, "inputs.csv")
        results = os.path.join(case_dir, "results.csv")
        docx = os.path.join(case_dir, "MO32_Crane_Compliance_Report.docx")
        if not os.path.isfile(inputs): continue
        try:
            df_in = pd.read_csv(inputs)
            vessel = str(df_in.get("Vessel Name").iloc[0]) if "Vessel Name" in df_in.columns and len(df_in) else ""
            imo = str(df_in.get("IMO").iloc[0]) if "IMO" in df_in.columns and len(df_in) else ""
            date_guess = d.replace("case_", "")
            rows.append({
                "case": d, "vessel": vessel, "imo": imo, "date": date_guess,
                "inputs": inputs, "results": results if os.path.isfile(results) else "",
                "docx": docx if os.path.isfile(docx) else ""
            })
        except Exception:
            continue

    def match(row):
        ok = True
        if q_name: ok = ok and (q_name.lower() in (row["vessel"] or "").lower())
        if q_imo: ok = ok and (q_imo.lower() in (row["imo"] or "").lower())
        return ok

    filtered = [r for r in rows if match(r)]
    if q_name or q_imo:
        st.write(f"Found {len(filtered)} matching cases.")
    else:
        st.write(f"Showing {len(filtered)} cases. Enter search terms to filter.")

    if not filtered:
        st.info("No matches in folders.")
        return

    for r in filtered:
        with st.container():
            c1, c2, c3, c4 = st.columns([3,2,2,2])
            c1.markdown(f"**{r['vessel'] or '(No Vessel Name)'}**  \nIMO: {r['imo'] or '-'}")
            c2.markdown(f"**Case:** {r['case']}")
            c3.markdown(f"**Saved:** {r['date']}")
            if r["docx"] and os.path.isfile(r["docx"]):
                try:
                    data = open(r["docx"], "rb").read()
                    c4.download_button("Download DOCX", data, file_name=f"{r['case']}.docx", key=f"dl_{r['case']}")
                except Exception as e:
                    c4.write(f"(DOCX not readable: {e})")
            else:
                c4.write("(No DOCX)")
            st.divider()

# -------------------------
# Router
# -------------------------
with st.sidebar:
    st.markdown("## Pages")
    page = st.radio("Navigate", ["Vessel Inspection", "Search Vessels"], index=0, key="page_radio")

st.title("Crane Compliance - Checker")
if page == "Inspection":
    page_inspection()
else:
    page_search()
